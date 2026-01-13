"""
Document management API routes
文档管理API路由

CAASS v2.0 Phase 2: Added paragraph context baseline and whitelist extraction

Security Features:
- File size limit validation
- File extension validation
- MIME type validation (if python-magic is installed)
"""

import uuid
import io
import logging
from datetime import datetime
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import List, Dict, Optional
from pydantic import BaseModel

from src.db.database import get_db

# Try to import python-magic for MIME validation
# 尝试导入python-magic用于MIME验证
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logging.warning(
        "python-magic not installed. MIME type validation disabled. "
        "Install with: pip install python-magic (Linux/macOS) or pip install python-magic-bin (Windows)"
    )

# Try to import python-docx for additional .docx validation
# 尝试导入python-docx进行额外的.docx验证
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
from src.db.models import Document, Sentence as SentenceModel
from src.api.schemas import DocumentInfo
from src.core.preprocessor.segmenter import SentenceSegmenter
from src.core.preprocessor.reference_handler import ReferenceHandler
from src.core.preprocessor.paraphrase_detector import ParaphraseDetector
from src.core.analyzer.scorer import RiskScorer, ParagraphContext, calculate_context_baseline
from src.core.preprocessor.whitelist_extractor import WhitelistExtractor

router = APIRouter()
segmenter = SentenceSegmenter()
ref_handler = ReferenceHandler()
paraphrase_detector = ParaphraseDetector()
scorer = RiskScorer()
whitelist_extractor = WhitelistExtractor()


class TextUploadRequest(BaseModel):
    """
    Request body for text upload
    文本上传请求体
    """
    text: str


@router.get("/", response_model=List[DocumentInfo])
async def list_documents(
    db: AsyncSession = Depends(get_db)
):
    """
    List all documents
    列出所有文档
    """
    result = await db.execute(
        select(Document).order_by(Document.created_at.desc())
    )
    docs = result.scalars().all()

    documents = []
    for doc in docs:
        # Count sentences by risk level
        # 按风险等级统计句子
        sentences_result = await db.execute(
            select(SentenceModel).where(SentenceModel.document_id == doc.id)
        )
        sentences = sentences_result.scalars().all()

        high_count = sum(1 for s in sentences if s.risk_level == "high")
        medium_count = sum(1 for s in sentences if s.risk_level == "medium")
        low_count = sum(1 for s in sentences if s.risk_level == "low")

        documents.append(DocumentInfo(
            id=doc.id,
            filename=doc.filename,
            status=doc.status,
            total_sentences=len(sentences),
            high_risk_count=high_count,
            medium_risk_count=medium_count,
            low_risk_count=low_count,
            created_at=doc.created_at
        ))

    return documents


@router.post("/upload", response_model=DocumentInfo)
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db)
):
    """
    Upload a document for processing
    上传文档进行处理

    CAASS v2.0 Phase 2: Uses paragraph context baseline and whitelist extraction

    Security (安全措施):
    - File size limit: max_file_size_mb (default 5MB)
    - File type validation: .txt and .docx only
    """
    from src.config import get_settings
    settings = get_settings()

    # Security: Validate file type
    # 安全：验证文件类型
    filename = file.filename or "unnamed.txt"
    allowed_extensions = ['.txt', '.docx']
    file_ext = '.' + filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail={
                "error": "invalid_file_type",
                "message": f"Only {', '.join(allowed_extensions)} files are allowed",
                "message_zh": f"仅支持 {', '.join(allowed_extensions)} 格式文件"
            }
        )

    # Security: Stream-read file with size limit to prevent memory exhaustion DoS
    # 安全：流式读取文件并限制大小，防止内存耗尽DoS攻击
    max_size_bytes = settings.max_file_size_mb * 1024 * 1024
    chunk_size = 64 * 1024  # 64KB chunks
    content_chunks = []
    total_size = 0

    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        total_size += len(chunk)

        # CRITICAL: Check size BEFORE storing in memory
        # 关键：在存入内存之前检查大小
        if total_size > max_size_bytes:
            # Immediately stop reading and reject
            # 立即停止读取并拒绝
            raise HTTPException(
                status_code=413,
                detail={
                    "error": "file_too_large",
                    "message": f"File size exceeds {settings.max_file_size_mb}MB limit",
                    "message_zh": f"文件大小超过 {settings.max_file_size_mb}MB 限制"
                }
            )

        content_chunks.append(chunk)

    # Combine chunks into final content
    # 将分块组合为最终内容
    content = b''.join(content_chunks)

    # Security: MIME type validation (if python-magic is available)
    # 安全：MIME类型验证（如果python-magic可用）
    if MAGIC_AVAILABLE:
        mime_type = magic.from_buffer(content, mime=True)
        allowed_mimes = {
            'text/plain',  # .txt files
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx files
            'application/octet-stream',  # Sometimes returned for binary files, will verify extension
        }

        # Allow application/octet-stream only for .docx files
        # 仅对.docx文件允许application/octet-stream
        if mime_type == 'application/octet-stream' and file_ext != '.docx':
            raise HTTPException(
                status_code=400,
                detail={
                    "error": "invalid_file_content",
                    "message": f"File content does not match expected type. Detected: {mime_type}",
                    "message_zh": f"文件内容与预期类型不匹配。检测到: {mime_type}"
                }
            )
        elif mime_type not in allowed_mimes:
            raise HTTPException(
                status_code=400,
                detail={
                    "error": "invalid_file_content",
                    "message": f"Invalid file content type. Detected: {mime_type}",
                    "message_zh": f"文件内容类型无效。检测到: {mime_type}"
                }
            )

    # Security: Validate .docx structure (if python-docx is available)
    # 安全：验证.docx结构（如果python-docx可用）
    if file_ext == '.docx' and DOCX_AVAILABLE:
        try:
            # Try to parse the docx file to ensure it's valid
            # 尝试解析docx文件以确保其有效
            DocxDocument(io.BytesIO(content))
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail={
                    "error": "invalid_docx",
                    "message": f"Invalid or corrupted DOCX file: {str(e)}",
                    "message_zh": f"无效或损坏的DOCX文件: {str(e)}"
                }
            )

    # Extract text based on file type
    # 根据文件类型提取文本
    if file_ext == '.docx':
        # For .docx files, extract text using python-docx to avoid XML/metadata
        # 对于.docx文件，使用python-docx提取文本以避免XML/元数据
        if not DOCX_AVAILABLE:
            raise HTTPException(
                status_code=500,
                detail={
                    "error": "docx_library_missing",
                    "message": "python-docx library is required to process .docx files",
                    "message_zh": "处理.docx文件需要python-docx库"
                }
            )

        try:
            # Parse .docx and extract only paragraph text (excludes headers, footers, comments, etc.)
            # 解析.docx并只提取段落文本（排除页眉、页脚、批注等）
            doc = DocxDocument(io.BytesIO(content))
            paragraphs = []
            for para in doc.paragraphs:
                # Only include non-empty paragraphs
                # 只包含非空段落
                para_text = para.text.strip()
                if para_text:
                    paragraphs.append(para_text)

            # Join paragraphs with double newline to preserve paragraph structure
            # 使用双换行连接段落以保留段落结构
            text = '\n\n'.join(paragraphs)

            if not text:
                raise HTTPException(
                    status_code=400,
                    detail={
                        "error": "empty_document",
                        "message": "No text content found in DOCX file",
                        "message_zh": "DOCX文件中未找到文本内容"
                    }
                )
        except Exception as e:
            # If already an HTTPException, re-raise it
            # 如果已经是HTTPException，重新抛出
            if isinstance(e, HTTPException):
                raise
            raise HTTPException(
                status_code=400,
                detail={
                    "error": "docx_parsing_failed",
                    "message": f"Failed to extract text from DOCX: {str(e)}",
                    "message_zh": f"从DOCX提取文本失败: {str(e)}"
                }
            )
    else:
        # For .txt files, decode as text
        # 对于.txt文件，解码为文本
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

    doc_id = str(uuid.uuid4())
    now = datetime.utcnow()

    # CAASS v2.0 Phase 2: Extract whitelist from document
    # CAASS v2.0 第二阶段：从文档提取白名单
    whitelist_result = whitelist_extractor.extract_from_document(text)
    whitelist_terms = whitelist_result.terms

    # Create document record with status "analyzing"
    # 创建文档记录，状态为"analyzing"
    doc = Document(
        id=doc_id,
        filename=filename,
        original_text=text,
        status="analyzing"
    )
    db.add(doc)
    await db.commit()

    # Extract reference section
    # 提取参考文献部分
    body_text, ref_section = ref_handler.extract(text)

    # CAASS v2.0 Phase 2: Segment text with paragraph info
    # CAASS v2.0 第二阶段：带段落信息分割文本
    # Only segment the body text
    sentences = segmenter.segment_with_paragraphs(body_text)

    # Build paragraph contexts for context baseline calculation
    # 构建段落上下文用于计算上下文基准
    paragraph_contexts = _build_paragraph_contexts(sentences)

    # Analyze each sentence and create records
    # 分析每个句子并创建记录
    high_count = 0
    medium_count = 0
    low_count = 0

    # Track processable sentence count for statistics
    # 跟踪可处理句子数量用于统计
    processable_count = 0

    for sent in sentences:
        sentence_id = str(uuid.uuid4())

        # Only analyze sentences that should be processed
        # 只分析应该处理的句子
        if sent.should_process:
            # Get paragraph context for this sentence
            # 获取该句子的段落上下文
            para_ctx = paragraph_contexts.get(sent.paragraph_index)

            # Analyze sentence risk (CAASS v2.0 Phase 2)
            # 分析句子风险（CAASS v2.0 第二阶段）
            analysis = scorer.analyze(
                sent.text,
                tone_level=4,
                whitelist=whitelist_terms,
                paragraph_context=para_ctx
            )
            
            # Detect paraphrase
            # 检测释义
            para_info = paraphrase_detector.detect(sent.text)
            
            processable_count += 1

            # Create sentence record with analysis
            # 创建带分析的句子记录
            sentence_record = SentenceModel(
                id=sentence_id,
                document_id=doc_id,
                index=sent.index,
                original_text=sent.text,
                content_type=sent.content_type,
                should_process=True,
                risk_score=analysis.risk_score,
                risk_level=analysis.risk_level,
                analysis_json={
                    "ppl": analysis.ppl,
                    "ppl_risk": analysis.ppl_risk,
                    "fingerprint_density": analysis.fingerprint_density,
                    "turnitin_score": analysis.turnitin_score,
                    "gptzero_score": analysis.gptzero_score,
                    "paragraph_index": sent.paragraph_index,
                    "context_baseline": para_ctx.baseline if para_ctx else 0,
                    # Phase 2: Enhanced metrics
                    "burstiness_value": getattr(analysis, 'burstiness_value', 0.0),
                    "burstiness_risk": getattr(analysis, 'burstiness_risk', 'unknown'),
                    "connector_count": getattr(analysis, 'connector_count', 0),
                    "connector_word": getattr(analysis, 'connector_match', None).connector if getattr(analysis, 'connector_match', None) else None,
                    # Paraphrase info
                    "is_paraphrase": para_info.is_paraphrase
                }
            )

            # Count risk levels only for processable sentences
            # 只统计可处理句子的风险等级
            if analysis.risk_level == "high":
                high_count += 1
            elif analysis.risk_level == "medium":
                medium_count += 1
            else:
                low_count += 1
        else:
            # Create sentence record without analysis (filtered content)
            # 创建不带分析的句子记录（过滤内容）
            sentence_record = SentenceModel(
                id=sentence_id,
                document_id=doc_id,
                index=sent.index,
                original_text=sent.text,
                content_type=sent.content_type,
                should_process=False,
                risk_score=0,
                risk_level="safe",
                analysis_json={
                    "filtered": True,
                    "reason": f"Content type: {sent.content_type}",
                    "paragraph_index": sent.paragraph_index
                }
            )

        db.add(sentence_record)

    # Add reference section as a special sentence if it exists
    # 如果存在参考文献部分，将其作为特殊句子添加
    if ref_section:
        ref_sentence_id = str(uuid.uuid4())
        ref_sentence = SentenceModel(
            id=ref_sentence_id,
            document_id=doc_id,
            index=len(sentences),  # Append at the end
            original_text=ref_section,
            content_type="reference_section",
            should_process=False,
            risk_score=0,
            risk_level="safe",
            analysis_json={
                "filtered": True,
                "reason": "Reference Section",
                "is_reference_blob": True
            }
        )
        db.add(ref_sentence)

    # Update document status to "ready" and store whitelist
    # 更新文档状态为"ready"并存储白名单
    result = await db.execute(
        select(Document).where(Document.id == doc_id)
    )
    doc_to_update = result.scalar_one()
    doc_to_update.status = "ready"
    await db.commit()

    return DocumentInfo(
        id=doc_id,
        filename=file.filename or "unnamed.txt",
        status="ready",
        total_sentences=len(sentences),
        high_risk_count=high_count,
        medium_risk_count=medium_count,
        low_risk_count=low_count,
        created_at=now
    )


def _build_paragraph_contexts(sentences) -> Dict[int, ParagraphContext]:
    """
    Build paragraph contexts from sentences
    从句子构建段落上下文

    Groups sentences by paragraph_index and creates ParagraphContext for each.
    按 paragraph_index 分组句子并为每个创建 ParagraphContext。
    """
    # Group sentences by paragraph
    # 按段落分组句子
    paragraphs: Dict[int, List[str]] = {}
    for sent in sentences:
        para_idx = sent.paragraph_index or 0
        if para_idx not in paragraphs:
            paragraphs[para_idx] = []
        if sent.should_process:
            paragraphs[para_idx].append(sent.text)

    # Build context for each paragraph
    # 为每个段落构建上下文
    contexts = {}
    for para_idx, sent_texts in paragraphs.items():
        if sent_texts:
            contexts[para_idx] = ParagraphContext.from_sentences(sent_texts)

    return contexts


@router.get("/{document_id}", response_model=DocumentInfo)
async def get_document(
    document_id: str,
    db: AsyncSession = Depends(get_db)
):
    """
    Get document information
    获取文档信息
    """
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Count sentences by risk level
    # 按风险等级统计句子
    sentences_result = await db.execute(
        select(SentenceModel).where(SentenceModel.document_id == document_id)
    )
    sentences = sentences_result.scalars().all()

    high_count = sum(1 for s in sentences if s.risk_level == "high")
    medium_count = sum(1 for s in sentences if s.risk_level == "medium")
    low_count = sum(1 for s in sentences if s.risk_level == "low")

    return DocumentInfo(
        id=doc.id,
        filename=doc.filename,
        status=doc.status,
        total_sentences=len(sentences),
        high_risk_count=high_count,
        medium_risk_count=medium_count,
        low_risk_count=low_count,
        created_at=doc.created_at,
        original_text=doc.original_text  # Include text for 5-layer analysis
    )


@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    db: AsyncSession = Depends(get_db)
):
    """
    Delete a document
    删除文档
    """
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    await db.delete(doc)
    await db.commit()

    return {"message": "Document deleted successfully"}


@router.post("/new/text")
async def upload_text(
    request: TextUploadRequest,
    db: AsyncSession = Depends(get_db)
):
    """
    Upload text directly (alternative to file upload)
    直接上传文本（文件上传的替代方案）

    CAASS v2.0 Phase 2: Uses paragraph context baseline and whitelist extraction

    Note: Text is received in request body to avoid URL length limits (431 error)
    注意：文本通过请求体接收以避免URL长度限制（431错误）
    """
    text = request.text
    doc_id = str(uuid.uuid4())

    # CAASS v2.0 Phase 2: Extract whitelist from text
    # CAASS v2.0 第二阶段：从文本提取白名单
    whitelist_result = whitelist_extractor.extract_from_document(text)
    whitelist_terms = whitelist_result.terms

    # Create document with status "analyzing"
    # 创建文档，状态为"analyzing"
    doc = Document(
        id=doc_id,
        filename="direct_input.txt",
        original_text=text,
        status="analyzing"
    )
    db.add(doc)
    await db.commit()

    # Extract reference section
    # 提取参考文献部分
    body_text, ref_section = ref_handler.extract(text)

    # CAASS v2.0 Phase 2: Segment text with paragraph info
    # CAASS v2.0 第二阶段：带段落信息分割文本
    sentences = segmenter.segment_with_paragraphs(body_text)

    # Build paragraph contexts for context baseline calculation
    # 构建段落上下文用于计算上下文基准
    paragraph_contexts = _build_paragraph_contexts(sentences)

    high_count = 0
    medium_count = 0
    low_count = 0

    for sent in sentences:
        sentence_id = str(uuid.uuid4())

        if sent.should_process:
            # Get paragraph context for this sentence
            # 获取该句子的段落上下文
            para_ctx = paragraph_contexts.get(sent.paragraph_index)

            # CAASS v2.0 Phase 2: Use paragraph context and whitelist
            # CAASS v2.0 第二阶段：使用段落上下文和白名单
            analysis = scorer.analyze(
                sent.text,
                tone_level=4,
                whitelist=whitelist_terms,
                paragraph_context=para_ctx
            )
            
            # Detect paraphrase
            # 检测释义
            para_info = paraphrase_detector.detect(sent.text)
            
            sentence_record = SentenceModel(
                id=sentence_id,
                document_id=doc_id,
                index=sent.index,
                original_text=sent.text,
                content_type=sent.content_type,
                should_process=True,
                risk_score=analysis.risk_score,
                risk_level=analysis.risk_level,
                analysis_json={
                    "ppl": analysis.ppl,
                    "ppl_risk": analysis.ppl_risk,
                    "fingerprint_density": analysis.fingerprint_density,
                    "paragraph_index": sent.paragraph_index,
                    "context_baseline": para_ctx.baseline if para_ctx else 0,
                    # Phase 2: Enhanced metrics
                    "burstiness_value": getattr(analysis, 'burstiness_value', 0.0),
                    "burstiness_risk": getattr(analysis, 'burstiness_risk', 'unknown'),
                    "connector_count": getattr(analysis, 'connector_count', 0),
                    "connector_word": getattr(analysis, 'connector_match', None).connector if getattr(analysis, 'connector_match', None) else None,
                    # Paraphrase info
                    "is_paraphrase": para_info.is_paraphrase
                }
            )

            if analysis.risk_level == "high":
                high_count += 1
            elif analysis.risk_level == "medium":
                medium_count += 1
            else:
                low_count += 1
        else:
            sentence_record = SentenceModel(
                id=sentence_id,
                document_id=doc_id,
                index=sent.index,
                original_text=sent.text,
                content_type=sent.content_type,
                should_process=False,
                risk_score=0,
                risk_level="safe",
                analysis_json={
                    "filtered": True,
                    "reason": f"Content type: {sent.content_type}",
                    "paragraph_index": sent.paragraph_index
                }
            )

        db.add(sentence_record)

    # Add reference section as a special sentence if it exists
    # 如果存在参考文献部分，将其作为特殊句子添加
    if ref_section:
        ref_sentence_id = str(uuid.uuid4())
        ref_sentence = SentenceModel(
            id=ref_sentence_id,
            document_id=doc_id,
            index=len(sentences),
            original_text=ref_section,
            content_type="reference_section",
            should_process=False,
            risk_score=0,
            risk_level="safe",
            analysis_json={
                "filtered": True,
                "reason": "Reference Section",
                "is_reference_blob": True
            }
        )
        db.add(ref_sentence)

    # Update document status to "ready"
    # 更新文档状态为"ready"
    result = await db.execute(
        select(Document).where(Document.id == doc_id)
    )
    doc_to_update = result.scalar_one()
    doc_to_update.status = "ready"
    await db.commit()

    return {
        "id": doc_id,
        "status": "ready",
        "total_sentences": len(sentences),
        "high_risk_count": high_count,
        "medium_risk_count": medium_count,
        "low_risk_count": low_count
    }
